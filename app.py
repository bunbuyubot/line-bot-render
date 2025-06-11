from flask import Flask, request, abort
from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError
from linebot.models import MessageEvent, TextMessage, TextSendMessage
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)

# 🔽 あなたのLINEチャネル情報を貼ってください
LINE_CHANNEL_ACCESS_TOKEN='JLHxkWqodOnZUYjdekyGfVPGecu8/QbV3v9b3/9v3QUVBt1e2VVa9iYEtjlZfyryyZ94VzBEFVjDVHhiifQybVHEgd/9G1YTyXNtpRYKYlS84prGTlQ9OEtjbYRQ0i+1Ew/LYVBKL/gOO8o28qXUNgdB04t89/1O/w1cDnyilFU='
LINE_CHANNEL_SECRET ='c30c9f9ecce29c412c0f912f56609edd'

line_bot_api = LineBotApi(LINE_CHANNEL_ACCESS_TOKEN)
handler = WebhookHandler(LINE_CHANNEL_SECRET)

SAVE_DIR = './reports'
os.makedirs(SAVE_DIR, exist_ok=True)

@app.route("/webhook", methods=['POST'])
def webhook():
    signature = request.headers['X-Line-Signature']
    body = request.get_data(as_text=True)

    try:
        handler.handle(body, signature)
    except InvalidSignatureError:
        abort(400)

    return 'OK'

@handler.add(MessageEvent, message=TextMessage)
def handle_message(event):
    print("🟢 handle_message() が呼ばれました")
    user_message = event.message.text
    user_id = event.source.user_id

    # Botから返信
    line_bot_api.reply_message(
        event.reply_token,
        TextSendMessage(text="メッセージを報告書に保存しました！")
    )

    # Wordファイル保存
    save_to_word(user_message, user_id)
    


def save_to_word(text, user_id):
    from datetime import datetime
    from docx import Document
    import os

    now = datetime.now()
    SAVE_DIR = '/tmp/reports'
    os.makedirs(SAVE_DIR, exist_ok=True)

    filename = f"report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(SAVE_DIR, filename)

    print(f"📄 Wordファイル作成準備中: {filepath}")
    upload_to_drive(filepath, filename)
    
    try:
        doc = Document()
        doc.add_heading("LINE報告書", level=1)
        doc.add_paragraph(f"日時: {now.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"ユーザーID: {user_id}")
        doc.add_paragraph("内容:")
        doc.add_paragraph(text)
        doc.save(filepath)
        print(f"✅ Wordファイルを保存しました: {filepath}")
    except Exception as e:
        print(f"❌ Word保存中にエラー発生: {e}")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)

def upload_to_drive(filepath, filename):
    print(f"🚀 Google Drive へアップロード開始: {filepath}")  # ←追加①

    import json
    import os
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    credentials_info = json.loads(os.environ.get("GOOGLE_CREDENTIALS_JSON"))
    credentials = service_account.Credentials.from_service_account_info(
        credentials_info,
        scopes=["https://www.googleapis.com/auth/drive.file"]
    )

    service = build("drive", "v3", credentials=credentials)

    file_metadata = {
    "name": filename,
    "parents": ["1TzWC2J5JBJXx4nr7Uu5nSHg-HUnQvh0v"]  
}


    uploaded = service.files().create(body=file_metadata, media_body=media, fields="id").execute()
    print(f"✅ Google Drive にアップロードされました (ID: {uploaded.get('id')})")

    print(f"✅ アップロード完了: https://drive.google.com/drive/u/0/my-drive")  # ←追加②





